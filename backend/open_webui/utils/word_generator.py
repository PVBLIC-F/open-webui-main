"""
Professional Word Document Generator for Chat Conversations using python-docx.

This module provides high-quality Word export with:
- Proper markdown rendering (headers, code blocks, lists, tables)
- Professional styling (color-coded messages, headers)
- Editable Word documents (users can edit after export)
- Small file sizes (text-based, not images)
- Fast generation (<2s for 100 messages)
"""

import re
import logging
from datetime import datetime
from io import BytesIO
from typing import Dict, Any, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from open_webui.models.chats import ChatTitleMessagesForm
from open_webui.env import SRC_LOG_LEVELS

log = logging.getLogger(__name__)
log.setLevel(SRC_LOG_LEVELS["MAIN"])


class ChatWordGenerator:
    """
    Modern Word document generator for chat conversations using python-docx.

    Features:
    - Professional styling with color-coded messages
    - Proper markdown parsing (headers, code, lists, tables, bold, italic)
    - Editable Word documents
    - Color-coded messages (user=green, assistant=blue)
    - Small file sizes (text-based)
    - Fast generation (<2s for 100 messages)

    Performance Optimizations:
    - Compiled regex patterns (class-level)
    - Efficient placeholder-based inline parsing
    - Constants for colors, spacing, and font sizes
    - Helper methods for common operations
    """

    # Color scheme (RGB values for Word)
    COLOR_USER = RGBColor(45, 90, 45)  # Green
    COLOR_ASSISTANT = RGBColor(45, 74, 90)  # Blue
    COLOR_SYSTEM = RGBColor(90, 45, 45)  # Red
    COLOR_CODE_BG = RGBColor(245, 245, 245)  # Light gray
    COLOR_CODE_TEXT = RGBColor(214, 51, 132)  # Pink/Magenta
    COLOR_HEADER = RGBColor(26, 26, 26)  # Dark gray
    COLOR_TIMESTAMP = RGBColor(128, 128, 128)  # Gray
    COLOR_WHITE = RGBColor(255, 255, 255)  # White

    # Spacing constants (in points)
    SPACING_SMALL = Pt(3)
    SPACING_MEDIUM = Pt(6)
    SPACING_LARGE = Pt(12)

    # Font sizes (in points)
    FONT_SIZE_CODE = Pt(9)
    FONT_SIZE_TIMESTAMP = Pt(9)
    FONT_SIZE_METADATA = Pt(10)
    FONT_SIZE_HEADER = Pt(12)

    # Indentation (in inches)
    INDENT_CODE = Inches(0.25)
    INDENT_QUOTE = Inches(0.5)

    # Table header background color (hex)
    TABLE_HEADER_BG = "2d4a5a"

    # Compiled regex patterns (performance optimization)
    _REGEX_CODE_BLOCK = re.compile(r"`([^`]+)`")
    _REGEX_BOLD = re.compile(r"\*\*(.+?)\*\*|__(.+?)__")
    _REGEX_ITALIC = re.compile(r"(?<!\*)\*(?!\*)([^*]+?)(?<!\*)\*(?!\*)|(?<!_)_(?!_)([^_]+?)(?<!_)_(?!_)")
    _REGEX_PLACEHOLDER = re.compile(r"(__(?:CODE|BOLD|ITALIC)_\d+__)")
    _REGEX_REASONING_BLOCK = re.compile(
        r'<details\s+type="reasoning"[^>]*>.*?</details>',
        flags=re.DOTALL | re.IGNORECASE
    )
    _REGEX_REASONING_SUMMARY = re.compile(
        r"<summary>Thought for \d+ seconds?</summary>",
        flags=re.IGNORECASE
    )
    _REGEX_HEADER = re.compile(r"^(#{1,6})\s+(.+)$")
    _REGEX_ORDERED_LIST = re.compile(r"^\d+\.\s+")
    _REGEX_ORDERED_LIST_ITEM = re.compile(r"^\d+\.\s+(.+)$")
    _REGEX_EMOJI = re.compile(
        "["
        "\U0001F300-\U0001F9FF"  # Miscellaneous Symbols and Pictographs
        "\U0001FA00-\U0001FAFF"  # Symbols and Pictographs Extended-A
        "\U00002600-\U000027BF"  # Miscellaneous Symbols
        "\U0001F900-\U0001F9FF"  # Supplemental Symbols and Pictographs
        "\U00002700-\U000027BF"  # Dingbats
        "\U0001F600-\U0001F64F"  # Emoticons
        "\U0001F680-\U0001F6FF"  # Transport and Map
        "\U00002600-\U000026FF"  # Miscellaneous symbols
        "\U0001F1E0-\U0001F1FF"  # Flags
        "]+",
        flags=re.UNICODE,
    )
    _REGEX_WHITESPACE = re.compile(r"\s+")

    def __init__(self, form_data: ChatTitleMessagesForm):
        """
        Initialize Word document generator.

        Args:
            form_data: Chat data with title and messages
        """
        self.form_data = form_data
        self.doc = Document()
        self._setup_document()

    def _setup_document(self):
        """Configure document settings (margins, etc.)."""
        # Set margins (1 inch)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _format_timestamp(self, timestamp: Any) -> Optional[str]:
        """
        Format timestamp to human-readable string.

        Args:
            timestamp: Unix timestamp (int or float) or ISO string

        Returns:
            Formatted timestamp string or None
        """
        try:
            if isinstance(timestamp, (int, float)):
                dt = datetime.fromtimestamp(timestamp)
            elif isinstance(timestamp, str):
                dt = datetime.fromisoformat(timestamp.replace("Z", "+00:00"))
            else:
                return None

            return dt.strftime("%B %d, %Y at %H:%M")
        except Exception as e:
            log.warning(f"Error formatting timestamp {timestamp}: {e}")
            return None

    def _escape_html(self, text: str) -> str:
        """
        Unescape HTML entities from LLM responses.

        Args:
            text: Text with HTML entities

        Returns:
            Text with entities properly handled
        """
        import html

        # Unescape HTML entities (e.g., &amp; -> &)
        return html.unescape(text)

    def _remove_emojis(self, text: str) -> str:
        """
        Remove emoji characters from text for clean Word output.

        Emojis don't render well in Word documents (require special fonts).
        Removing them creates cleaner, more professional documents.

        Args:
            text: Text that may contain emojis

        Returns:
            Text with emojis removed, extra whitespace cleaned
        """
        # Remove emojis using compiled pattern
        text = self._REGEX_EMOJI.sub("", text)

        # Clean up extra whitespace left by emoji removal
        text = self._REGEX_WHITESPACE.sub(" ", text).strip()

        return text

    def _filter_reasoning_blocks(self, content: str) -> str:
        """
        Remove LLM reasoning blocks from content before Word generation.

        Filters out <details type="reasoning"> blocks that show internal
        LLM thinking process - users don't need to see this in Word docs.

        Args:
            content: Message content with potential reasoning blocks

        Returns:
            Content with reasoning blocks removed
        """
        # Remove reasoning blocks using compiled patterns
        content = self._REGEX_REASONING_BLOCK.sub("", content)
        content = self._REGEX_REASONING_SUMMARY.sub("", content)

        return content.strip()

    def _parse_markdown_inline(self, text: str) -> List[Dict[str, Any]]:
        """
        Parse inline markdown (bold, italic, inline code) to Word runs.

        Uses placeholder substitution to handle nested formatting correctly.
        Process in priority: code > bold > italic.
        Uses compiled regex patterns for performance.

        Args:
            text: Markdown text

        Returns:
            List of run dictionaries with text and formatting
        """
        # Escape HTML first (security)
        text = self._escape_html(text)

        # Remove emojis for clean Word output
        text = self._remove_emojis(text)

        # Early return if no formatting markers
        if "`" not in text and "**" not in text and "__" not in text and "*" not in text and "_" not in text:
            return [{"text": text}]

        runs = []
        replacements = {}  # placeholder -> (type, text)

        # Step 1: Replace code blocks with placeholders (highest priority)
        code_counter = 0
        def code_replacer(match):
            nonlocal code_counter
            placeholder = f"__CODE_{code_counter}__"
            replacements[placeholder] = ("code", match.group(1))
            code_counter += 1
            return placeholder
        text = self._REGEX_CODE_BLOCK.sub(code_replacer, text)

        # Step 2: Replace bold text with placeholders
        bold_counter = 0
        def bold_replacer(match):
            nonlocal bold_counter
            placeholder = f"__BOLD_{bold_counter}__"
            bold_text = match.group(1) or match.group(2)
            replacements[placeholder] = ("bold", bold_text)
            bold_counter += 1
            return placeholder
        text = self._REGEX_BOLD.sub(bold_replacer, text)

        # Step 3: Replace italic text with placeholders (single * or _, not ** or __)
        italic_counter = 0
        def italic_replacer(match):
            nonlocal italic_counter
            placeholder = f"__ITALIC_{italic_counter}__"
            italic_text = match.group(1) or match.group(2)
            replacements[placeholder] = ("italic", italic_text)
            italic_counter += 1
            return placeholder
        text = self._REGEX_ITALIC.sub(italic_replacer, text)

        # Step 4: Split text by placeholders and build runs
        parts = self._REGEX_PLACEHOLDER.split(text)

        for part in parts:
            if part in replacements:
                placeholder_type, placeholder_text = replacements[part]
                if placeholder_type == "code":
                    runs.append(
                        {
                            "text": placeholder_text,
                            "font": "Courier New",
                            "color": self.COLOR_CODE_TEXT,
                        }
                    )
                elif placeholder_type == "bold":
                    runs.append({"text": placeholder_text, "bold": True})
                elif placeholder_type == "italic":
                    runs.append({"text": placeholder_text, "italic": True})
            elif part:
                runs.append({"text": part})

        # Fallback: if no runs created, return plain text
        if not runs:
            runs = [{"text": text}]

        return runs

    def _add_formatted_run(self, paragraph, run_dict: Dict[str, Any]):
        """
        Add a formatted run to a paragraph.

        Args:
            paragraph: Word paragraph object
            run_dict: Dictionary with text and formatting (bold, italic, color, font)
        """
        run = paragraph.add_run(run_dict["text"])

        if run_dict.get("bold"):
            run.bold = True
        if run_dict.get("italic"):
            run.italic = True
        if run_dict.get("color"):
            run.font.color.rgb = run_dict["color"]
        if run_dict.get("font"):
            run.font.name = run_dict["font"]
        if run_dict.get("size"):
            run.font.size = run_dict["size"]

    def _add_spacing_paragraph(self, space_after: Optional[Pt] = None):
        """
        Add a spacing paragraph (optimized helper).

        Args:
            space_after: Spacing after paragraph (default: SPACING_SMALL)
        """
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = space_after or self.SPACING_SMALL

    def _parse_markdown_table(self, table_lines: List[str]) -> Tuple[Optional[List[List[str]]], Optional[List[str]]]:
        """
        Parse markdown table syntax into table data and alignments.

        Args:
            table_lines: List of table lines (header, separator, rows)

        Returns:
            Tuple of (table_data, alignments) or (None, None) if invalid
        """
        try:
            if len(table_lines) < 2:
                return None, None

            # Parse header
            header_line = table_lines[0]
            header_cells = [cell.strip() for cell in header_line.split("|")[1:-1]]
            num_cols = len(header_cells)

            # Parse separator to determine alignments
            separator = table_lines[1].strip()
            separator_cells = [cell.strip() for cell in separator.split("|")[1:-1]]

            alignments = []
            for cell in separator_cells:
                if cell.startswith(":") and cell.endswith(":"):
                    alignments.append("CENTER")
                elif cell.endswith(":"):
                    alignments.append("RIGHT")
                else:
                    alignments.append("LEFT")

            # Normalize alignments to match column count
            while len(alignments) < num_cols:
                alignments.append("LEFT")
            alignments = alignments[:num_cols]

            # Build table data
            table_data = [header_cells]

            # Parse data rows
            for line in table_lines[2:]:
                cells = [cell.strip() for cell in line.split("|")[1:-1]]

                # Normalize column count
                while len(cells) < num_cols:
                    cells.append("")
                cells = cells[:num_cols]

                table_data.append(cells)

            return table_data, alignments

        except Exception as e:
            log.warning(f"Error parsing markdown table: {e}")
            return None, None

    def _parse_markdown_to_paragraphs(self, content: str):
        """
        Parse markdown content into Word paragraphs.

        Supports:
        - Headers (# ## ###)
        - Code blocks (```)
        - Lists (- or 1.)
        - Block quotes (>)
        - Tables (| Header | Header |)
        - Paragraphs

        Args:
            content: Markdown text
        """
        # Filter out LLM reasoning blocks first
        content = self._filter_reasoning_blocks(content)

        lines = content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]

            # Code block: ``` ... ```
            if line.strip().startswith("```"):
                code_lines = []
                i += 1  # Skip opening ```

                # Collect code until closing ```
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1

                # Create code block paragraph
                if code_lines:
                    code_text = "\n".join(code_lines)
                    p = self.doc.add_paragraph()
                    p.style = "No Spacing"

                    # Add code with monospace font
                    run = p.add_run(code_text)
                    run.font.name = "Courier New"
                    run.font.size = self.FONT_SIZE_CODE
                    run.font.color.rgb = RGBColor(0, 0, 0)

                    # Set paragraph formatting
                    p.paragraph_format.left_indent = self.INDENT_CODE
                    p.paragraph_format.space_before = self.SPACING_MEDIUM
                    p.paragraph_format.space_after = self.SPACING_MEDIUM

                i += 1  # Skip closing ```
                continue

            # Header: # Header
            if line.startswith("#"):
                match = self._REGEX_HEADER.match(line)
                if match:
                    level = len(match.group(1))
                    text = match.group(2)

                    # Create header paragraph
                    p = self.doc.add_heading(level=min(level, 6))
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Add formatted runs
                    runs = self._parse_markdown_inline(text)
                    for run_dict in runs:
                        self._add_formatted_run(p, run_dict)

                    i += 1
                    continue

            # Unordered list: - item
            if line.strip().startswith("- "):
                list_items = []

                while i < len(lines) and lines[i].strip().startswith("- "):
                    item_text = lines[i].strip()[2:]  # Remove "- "
                    list_items.append(item_text)
                    i += 1

                # Create list items
                for item_text in list_items:
                    p = self.doc.add_paragraph(style="List Bullet")
                    runs = self._parse_markdown_inline(item_text)
                    for run_dict in runs:
                        self._add_formatted_run(p, run_dict)

                continue

            # Ordered list: 1. item
            if self._REGEX_ORDERED_LIST.match(line.strip()):
                list_items = []

                while i < len(lines) and self._REGEX_ORDERED_LIST.match(lines[i].strip()):
                    match = self._REGEX_ORDERED_LIST_ITEM.match(lines[i].strip())
                    if match:
                        item_text = match.group(1)
                    else:
                        # Fallback if regex doesn't match (shouldn't happen)
                        item_text = lines[i].strip().split(".", 1)[1].strip() if "." in lines[i].strip() else lines[i].strip()
                    list_items.append(item_text)
                    i += 1

                # Create list items
                for item_text in list_items:
                    p = self.doc.add_paragraph(style="List Number")
                    runs = self._parse_markdown_inline(item_text)
                    for run_dict in runs:
                        self._add_formatted_run(p, run_dict)

                continue

            # Block quote: > text
            if line.strip().startswith(">"):
                quote_lines = []

                while i < len(lines) and lines[i].strip().startswith(">"):
                    quote_text = lines[i].strip()[1:].strip()  # Remove ">"
                    if quote_text:
                        quote_lines.append(quote_text)
                    i += 1

                if quote_lines:
                    quote_text = " ".join(quote_lines)
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = self.INDENT_QUOTE
                    p.paragraph_format.space_before = self.SPACING_MEDIUM
                    p.paragraph_format.space_after = self.SPACING_MEDIUM

                    # Add italic formatted text
                    runs = self._parse_markdown_inline(quote_text)
                    for run_dict in runs:
                        run_dict["italic"] = True
                        self._add_formatted_run(p, run_dict)

                continue

            # Table: | Header | Header |
            if "|" in line and line.strip().startswith("|"):
                table_lines = []

                # Collect all table lines
                while i < len(lines) and "|" in lines[i] and lines[i].strip():
                    table_lines.append(lines[i])
                    i += 1

                # Need at least header + separator
                if len(table_lines) >= 2:
                    # Validate second line is separator row
                    separator = table_lines[1].strip()
                    is_separator = all(c in "|-: " for c in separator)

                    if not is_separator:
                        # Not a valid table, treat as regular text
                        for table_line in table_lines:
                            p = self.doc.add_paragraph()
                            runs = self._parse_markdown_inline(table_line)
                            for run_dict in runs:
                                self._add_formatted_run(p, run_dict)
                        continue

                    try:
                        # Parse table structure
                        table_data, alignments = self._parse_markdown_table(table_lines)

                        if table_data and len(table_data) > 0:
                            # Create Word table
                            table = self.doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                            table.alignment = WD_TABLE_ALIGNMENT.LEFT

                            # Populate table
                            for row_idx, row_data in enumerate(table_data):
                                for col_idx, cell_text in enumerate(row_data):
                                    cell = table.rows[row_idx].cells[col_idx]
                                    cell_para = cell.paragraphs[0]

                                    # Clear default paragraph
                                    cell_para.clear()

                                    # Add formatted text
                                    runs = self._parse_markdown_inline(cell_text)
                                    for run_dict in runs:
                                        self._add_formatted_run(cell_para, run_dict)

                                    # Set alignment
                                    if alignments and col_idx < len(alignments):
                                        align = alignments[col_idx]
                                        if align == "CENTER":
                                            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        elif align == "RIGHT":
                                            cell_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                        else:
                                            cell_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                                    # Header row styling (each cell needs its own shading element)
                                    if row_idx == 0:
                                        for run in cell_para.runs:
                                            run.bold = True
                                            run.font.color.rgb = self.COLOR_WHITE
                                        # Create shading element (using constant for color)
                                        cell_shading = OxmlElement(qn("w:shd"))
                                        cell_shading.set(qn("w:fill"), self.TABLE_HEADER_BG)
                                        cell_para._element.get_or_add_pPr().append(cell_shading)

                            # Add spacing after table
                            self._add_spacing_paragraph(self.SPACING_MEDIUM)

                    except Exception as e:
                        log.warning(f"Error rendering table: {e}")
                        # Fallback: render as plain text
                        for table_line in table_lines:
                            p = self.doc.add_paragraph()
                            runs = self._parse_markdown_inline(table_line)
                            for run_dict in runs:
                                self._add_formatted_run(p, run_dict)

                continue

            # Empty line
            if not line.strip():
                self._add_spacing_paragraph()
                i += 1
                continue

            # Regular paragraph
            if line.strip():
                p = self.doc.add_paragraph()
                runs = self._parse_markdown_inline(line)
                for run_dict in runs:
                    self._add_formatted_run(p, run_dict)

            i += 1

    def _build_message_paragraphs(self, message: Dict[str, Any]):
        """
        Build Word paragraphs for a single message.

        Args:
            message: Message dict with role, content, timestamp, model
        """
        role = message.get("role", "user").lower()
        content = message.get("content", "")
        timestamp = message.get("timestamp")
        model = message.get("model", "")

        # Determine header style based on role
        if role == "user":
            header_color = self.COLOR_USER
            role_display = "User"
        elif role == "assistant":
            header_color = self.COLOR_ASSISTANT
            role_display = "Assistant"
            if model:
                role_display += f" ({model})"
        else:
            header_color = self.COLOR_SYSTEM
            role_display = role.title()

        # Message header (remove emojis for clean output)
        clean_role_display = self._remove_emojis(role_display)
        p = self.doc.add_heading(level=3)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add colored header
        run = p.add_run(clean_role_display)
        run.font.color.rgb = header_color
        run.font.size = self.FONT_SIZE_HEADER
        run.bold = True

        # Timestamp (if available)
        if timestamp:
            timestamp_str = self._format_timestamp(timestamp)
            if timestamp_str:
                p = self.doc.add_paragraph()
                run = p.add_run(timestamp_str)
                run.italic = True
                run.font.size = self.FONT_SIZE_TIMESTAMP
                run.font.color.rgb = self.COLOR_TIMESTAMP

        # Handle empty content
        if not content or not content.strip():
            p = self.doc.add_paragraph()
            run = p.add_run("[Empty message]")
            run.italic = True
            run.font.color.rgb = self.COLOR_TIMESTAMP
            return

        # Parse markdown content
        self._parse_markdown_to_paragraphs(content)

        # Add spacing after message
        self._add_spacing_paragraph(self.SPACING_LARGE)

    def generate_chat_docx(self) -> bytes:
        """
        Generate Word document from chat data.

        Returns:
            Word document as bytes
        """
        try:
            # Title page
            clean_title = self._remove_emojis(self.form_data.title)
            title = self.doc.add_heading(clean_title, level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add spacing
            self._add_spacing_paragraph(self.SPACING_MEDIUM)

            # Chat metadata
            metadata_text = f"Exported on {datetime.now().strftime('%B %d, %Y at %H:%M')}"
            p = self.doc.add_paragraph()
            run = p.add_run(metadata_text)
            run.italic = True
            run.font.size = self.FONT_SIZE_METADATA
            run.font.color.rgb = self.COLOR_TIMESTAMP
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Message count
            msg_count = len(self.form_data.messages)
            count_text = f"{msg_count} message{'s' if msg_count != 1 else ''}"
            p = self.doc.add_paragraph()
            run = p.add_run(count_text)
            run.italic = True
            run.font.size = self.FONT_SIZE_METADATA
            run.font.color.rgb = self.COLOR_TIMESTAMP
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add spacing
            self._add_spacing_paragraph(self.SPACING_LARGE)

            # Add all messages
            log.info(f"Generating Word doc for chat '{self.form_data.title}' with {msg_count} messages")

            for idx, message in enumerate(self.form_data.messages):
                try:
                    self._build_message_paragraphs(message)

                except Exception as e:
                    log.error(f"Error processing message {idx}: {e}")
                    # Add error message to doc instead of failing
                    p = self.doc.add_paragraph()
                    run = p.add_run(f"[Error rendering message {idx + 1}]")
                    run.italic = True
                    run.font.color.rgb = RGBColor(255, 0, 0)

            # Save to bytes
            buffer = BytesIO()
            self.doc.save(buffer)
            docx_bytes = buffer.getvalue()
            buffer.close()

            log.info(f"Successfully generated Word doc: {len(docx_bytes) / 1024:.1f} KB")

            return docx_bytes

        except Exception as e:
            log.exception(f"Error generating Word doc: {e}")
            raise


# Backward-compatible wrapper
class WordGenerator(ChatWordGenerator):
    """
    Backward-compatible wrapper for ChatWordGenerator.

    Deprecated: Use ChatWordGenerator directly.
    """

    def __init__(self, form_data: ChatTitleMessagesForm):
        log.warning("WordGenerator is deprecated. Use ChatWordGenerator instead.")
        super().__init__(form_data)
